# """
# Resume Parser — Zero API calls at upload time.

# Text extraction  : pdfplumber / PyMuPDF / python-docx
# Skill extraction : section-aware tokenization — skills are extracted ONLY
#                    from skill/project/experience sections, NOT from name,
#                    address, or contact lines. This prevents personal-info
#                    tokens (e.g. "faisal khan", "address labeling") from
#                    polluting the skill bag.
# Experience years : date-range scan, education-section excluded
# """

# import hashlib
# import re
# import pdfplumber
# import fitz  # PyMuPDF
# from pathlib import Path
# from docx import Document


# # ── Section headers ───────────────────────────────────────────────────────────
# SECTION_HEADERS = {
#     "experience":     {
#         "work experience", "professional experience", "employment history",
#         "experience", "work history", "career history",
#     },
#     "education":      {
#         "education", "academic background", "qualifications",
#         "degrees", "academic qualifications",
#     },
#     "skills":         {
#         "skills", "technical skills", "core competencies", "technologies",
#         "tech stack", "expertise", "key skills", "competencies",
#         "tools & technologies", "tools and technologies",
#         "tools", "software", "platforms", "tech tools", "applications",
#     },
#     "projects":       {
#         "projects", "personal projects", "key projects",
#         "portfolio", "selected projects", "project experience",
#     },
#     "certifications": {
#         "certifications", "certificates", "awards",
#         "achievements", "licenses", "credentials",
#     },
#     "summary":        {
#         "summary", "objective", "profile", "about me",
#         "overview", "professional summary", "career objective",
#         "executive summary",
#     },
# }

# # Stop-words to drop from skill tokens
# STOP_WORDS = {
#     "and", "or", "the", "a", "an", "in", "on", "at", "to", "for",
#     "of", "with", "is", "are", "was", "were", "be", "been", "have",
#     "has", "had", "do", "did", "will", "would", "could", "should",
#     "may", "might", "i", "me", "my", "we", "our", "you", "your",
#     "he", "she", "it", "they", "their", "this", "that", "these",
#     "those", "not", "no", "nor", "but", "so", "yet", "both", "either",
#     "as", "if", "then", "than", "by", "from", "up", "about", "into",
#     "through", "during", "including", "used", "using", "worked",
#     "work", "working", "experience", "strong", "good", "knowledge",
#     "understanding", "ability", "skills", "skill", "years", "year",
#     "month", "months", "present", "current", "responsible", "team",
#     "also", "well", "new", "use", "develop", "development",
#     # ── NEW: names / address noise ────────────────────────────────────────────
#     "mr", "ms", "dr", "sir", "dear", "hello", "hi",
#     "street", "road", "lane", "avenue", "ave", "blvd", "st", "rd",
#     "city", "state", "country", "zip", "postcode", "uk", "us", "usa",
#     "phone", "mobile", "tel", "email", "linkedin", "github", "portfolio",
#     "page", "cv", "resume",
# }

# # ── Patterns that look like personal-info lines (skip entirely for skills) ────
# _PERSONAL_INFO_PREFIXES = (
#     "email", "phone", "mobile", "tel", "address", "linkedin",
#     "github", "twitter", "website", "portfolio", "location",
#     "nationality", "date of birth", "dob", "gender",
# )


# def _is_personal_info_line(line: str) -> bool:
#     low = line.lower().strip()
#     if any(low.startswith(p) for p in _PERSONAL_INFO_PREFIXES):
#         return True
#     # Line that looks like a postal address (contains digit + word pattern early)
#     if re.match(r"^\d+\s+\w", low):
#         return True
#     # Line that is just a name (2-4 capitalised words, no digits/punctuation)
#     words = low.split()
#     if 2 <= len(words) <= 4 and all(w.isalpha() for w in words):
#         return True
#     return False


# def file_hash(content: bytes) -> str:
#     return hashlib.sha256(content).hexdigest()[:16]


# # ── Text extractors ───────────────────────────────────────────────────────────
# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     import io
#     text = ""
#     try:
#         with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text()
#                 if t:
#                     text += t + "\n"
#         if text.strip():
#             return text
#     except Exception:
#         pass
#     try:
#         doc = fitz.open(stream=file_bytes, filetype="pdf")
#         for page in doc:
#             text += page.get_text() + "\n"
#     except Exception:
#         pass
#     return text


# def extract_text_from_docx(file_bytes: bytes) -> str:
#     import io
#     doc = Document(io.BytesIO(file_bytes))
#     return "\n".join(p.text for p in doc.paragraphs)


# def extract_text_from_txt(file_bytes: bytes) -> str:
#     for enc in ("utf-8", "latin-1", "cp1252"):
#         try:
#             return file_bytes.decode(enc)
#         except Exception:
#             continue
#     return ""


# def extract_text(file_bytes: bytes, filename: str) -> str:
#     ext = Path(filename).suffix.lower()
#     if ext == ".pdf":
#         return extract_text_from_pdf(file_bytes)
#     elif ext in (".docx", ".doc"):
#         return extract_text_from_docx(file_bytes)
#     elif ext == ".txt":
#         return extract_text_from_txt(file_bytes)
#     return ""


# # ── Heuristic helpers ─────────────────────────────────────────────────────────
# def extract_email(text: str) -> str:
#     for token in text.split():
#         token = token.strip("(),;<>\"'\n\r")
#         if "@" in token and "." in token.split("@")[-1] and len(token) > 5:
#             return token
#     return ""


# def extract_name(text: str) -> str:
#     lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
#     for line in lines[:6]:
#         if "@" in line or len(line) > 50:
#             continue
#         words = line.split()
#         if 2 <= len(words) <= 5:
#             if all(ch.isalpha() or ch in " .-'" for ch in line):
#                 return line
#     return lines[0][:40] if lines else "Unknown"


# def extract_experience_years(text: str, sections: dict = None) -> float:
#     import datetime
#     current_year = datetime.datetime.now().year
#     ranges = []
#     if sections is None:
#         sections = split_sections(text)
#     edu_text  = sections.get("education", "")
#     work_text = text.replace(edu_text, " ") if edu_text.strip() else text

#     year_pattern = re.compile(
#         r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
#         r'((?:19|20)\d{2})'
#         r'\s*(?:-|–|—|to)\s*'
#         r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
#         r'((?:19|20)\d{2}|present|current|now|till date|date)',
#         re.IGNORECASE
#     )

#     for match in year_pattern.finditer(work_text):
#         s       = int(match.group(2))
#         end_raw = match.group(4).strip().lower()
#         e       = current_year if end_raw in ("present", "current", "now", "till date", "date") else int(end_raw)
#         if not (1980 <= s <= current_year and s <= e <= current_year + 1):
#             continue
#         ranges.append((s, e))

#     if ranges:
#         ranges.sort()
#         merged = [list(ranges[0])]
#         for start, end in ranges[1:]:
#             if start <= merged[-1][1]:
#                 merged[-1][1] = max(merged[-1][1], end)
#             else:
#                 merged.append([start, end])
#         total = sum(e - s for s, e in merged)
#         return round(min(total, 40.0), 1)

#     fallback_pattern = re.compile(
#         r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?(?:\s+experience)?',
#         re.IGNORECASE
#     )
#     matches = fallback_pattern.findall(work_text)
#     if matches:
#         return round(min(max(int(m) for m in matches), 40.0), 1)
#     return 0.0


# def extract_skills_from_text(text: str, skill_sections_text: str = "") -> list:
#     """
#     KEY FIX: Extract skills ONLY from skill-relevant sections.

#     Priority order:
#       1. If we have dedicated skill/experience/project section text → use ONLY that.
#       2. Otherwise fall back to full text BUT skip personal-info lines.

#     This prevents name tokens, address tokens, and email prefixes from
#     becoming "skills" that then get fuzzy-matched against JD terms.

#     No hardcoded taxonomy — LLM does semantic matching at query time.
#     """
#     skills = []
#     seen   = set()

#     # Use skill-section text if available, else fall back to full text
#     source_text = skill_sections_text.strip() if skill_sections_text.strip() else text

#     for line in source_text.split("\n"):
#         line = line.strip()
#         if len(line) < 3 or len(line) > 120:
#             continue

#         # ── NEW: skip personal info lines entirely ────────────────────────────
#         if _is_personal_info_line(line):
#             continue

#         cleaned = ""
#         for ch in line:
#             if ch.isalnum() or ch in " .+#/-_":
#                 cleaned += ch
#             else:
#                 cleaned += " "

#         raw_tokens = [t.strip(".-_/") for t in cleaned.split() if t.strip(".-_/")]

#         for i, tok in enumerate(raw_tokens):
#             tok_lower = tok.lower()
#             # ── NEW: skip very short tokens and pure-digit tokens ─────────────
#             if len(tok_lower) < 3 or tok_lower in STOP_WORDS or tok_lower.isdigit():
#                 continue
#             if tok_lower not in seen:
#                 seen.add(tok_lower)
#                 skills.append(tok_lower)

#             # 2-gram — only if first word is not a stop-word
#             if i + 1 < len(raw_tokens):
#                 next_tok = raw_tokens[i + 1].lower()
#                 if next_tok in STOP_WORDS or next_tok.isdigit():
#                     continue
#                 bigram = (tok_lower + " " + next_tok).strip()
#                 if bigram not in seen and len(bigram) <= 40:
#                     parts = bigram.split()
#                     if not any(p in STOP_WORDS for p in parts):
#                         seen.add(bigram)
#                         skills.append(bigram)

#     return skills[:200]


# # ── Section splitter ──────────────────────────────────────────────────────────
# def split_sections(text: str) -> dict:
#     lines    = text.split("\n")
#     sections = {k: [] for k in ["full", "skills", "experience", "projects",
#                                  "education", "certifications", "summary"]}
#     current  = "full"

#     for line in lines:
#         stripped = line.strip().lower().rstrip(":").strip()
#         matched  = None
#         for section, headers in SECTION_HEADERS.items():
#             if stripped in headers:
#                 matched = section
#                 break
#         if matched:
#             current = matched
#         else:
#             sections[current].append(line)

#     return {k: "\n".join(v).strip() for k, v in sections.items()}


# # ── Main pipeline (0 API calls) ───────────────────────────────────────────────
# def parse_resume(file_bytes: bytes, filename: str) -> dict:
#     raw_text = extract_text(file_bytes, filename)
#     if not raw_text.strip():
#         return {}

#     sections  = split_sections(raw_text)
#     exp_years = extract_experience_years(raw_text, sections)

#     # ── FIX: build skill-section source text (exclude header + contact area) ──
#     # Combine skill-relevant sections; fall back to full text only if empty
#     skill_source_parts = [
#         sections.get("skills", ""),
#         sections.get("experience", ""),
#         sections.get("projects", ""),
#         sections.get("summary", ""),
#         sections.get("certifications", ""),
#     ]
#     skill_source = "\n".join(p for p in skill_source_parts if p.strip())

#     skills = extract_skills_from_text(raw_text, skill_sections_text=skill_source)

#     chunk_labels = {
#         "skills":         sections.get("skills") or " ".join(skills[:60]),
#         "experience":     sections.get("experience") or "",
#         "projects":       sections.get("projects") or "",
#         "education":      sections.get("education") or "",
#         "certifications": sections.get("certifications") or "",
#         "summary":        sections.get("summary") or "",
#         "full":           raw_text[:3000],
#     }
#     chunks = {label: text.strip() for label, text in chunk_labels.items() if text.strip()}

#     return {
#         "filename":         filename,
#         "file_hash":        file_hash(file_bytes),
#         "name":             extract_name(raw_text),
#         "email":            extract_email(raw_text),
#         "skills":           skills,
#         "experience_years": exp_years,
#         "raw_text":         raw_text,
#         "chunks":           chunks,
#     }


# this is final 
# """
# Resume Parser — Zero API calls at upload time.

# Text extraction  : pdfplumber / PyMuPDF / python-docx
# Skill extraction : section-aware tokenization — skills are extracted ONLY
#                    from skill/project/experience sections, NOT from name,
#                    address, or contact lines. This prevents personal-info
#                    tokens (e.g. "faisal khan", "address labeling") from
#                    polluting the skill bag.
# Experience years : date-range scan, education-section excluded
# """

# import hashlib
# import re
# import pdfplumber
# import fitz  # PyMuPDF
# from pathlib import Path
# from docx import Document


# # ── Section headers ───────────────────────────────────────────────────────────
# SECTION_HEADERS = {
#     "experience":     {
#         "work experience", "professional experience", "employment history",
#         "experience", "work history", "career history",
#     },
#     "education":      {
#         "education", "academic background", "qualifications",
#         "degrees", "academic qualifications",
#     },
#     "skills":         {
#         "skills", "technical skills", "core competencies", "technologies",
#         "tech stack", "expertise", "key skills", "competencies",
#         "tools & technologies", "tools and technologies",
#         "tools", "software", "platforms", "tech tools", "applications",
#     },
#     "projects":       {
#         "projects", "personal projects", "key projects",
#         "portfolio", "selected projects", "project experience",
#     },
#     "certifications": {
#         "certifications", "certificates", "awards",
#         "achievements", "licenses", "credentials",
#     },
#     "summary":        {
#         "summary", "objective", "profile", "about me",
#         "overview", "professional summary", "career objective",
#         "executive summary",
#     },
# }

# # Stop-words to drop from skill tokens
# STOP_WORDS = {
#     "and", "or", "the", "a", "an", "in", "on", "at", "to", "for",
#     "of", "with", "is", "are", "was", "were", "be", "been", "have",
#     "has", "had", "do", "did", "will", "would", "could", "should",
#     "may", "might", "i", "me", "my", "we", "our", "you", "your",
#     "he", "she", "it", "they", "their", "this", "that", "these",
#     "those", "not", "no", "nor", "but", "so", "yet", "both", "either",
#     "as", "if", "then", "than", "by", "from", "up", "about", "into",
#     "through", "during", "including", "used", "using", "worked",
#     "work", "working", "experience", "strong", "good", "knowledge",
#     "understanding", "ability", "skills", "skill", "years", "year",
#     "month", "months", "present", "current", "responsible", "team",
#     "also", "well", "new", "use", "develop", "development",
#     # ── NEW: names / address noise ────────────────────────────────────────────
#     "mr", "ms", "dr", "sir", "dear", "hello", "hi",
#     "street", "road", "lane", "avenue", "ave", "blvd", "st", "rd",
#     "city", "state", "country", "zip", "postcode", "uk", "us", "usa",
#     "phone", "mobile", "tel", "email", "linkedin", "github", "portfolio",
#     "page", "cv", "resume",
# }

# # ── Patterns that look like personal-info lines (skip entirely for skills) ────
# _PERSONAL_INFO_PREFIXES = (
#     "email", "phone", "mobile", "tel", "address", "linkedin",
#     "github", "twitter", "website", "portfolio", "location",
#     "nationality", "date of birth", "dob", "gender",
# )


# def _is_personal_info_line(line: str) -> bool:
#     low = line.lower().strip()
#     if any(low.startswith(p) for p in _PERSONAL_INFO_PREFIXES):
#         return True
#     # Line that looks like a postal address (contains digit + word pattern early)
#     if re.match(r"^\d+\s+\w", low):
#         return True
#     # Line that is just a name (2-4 capitalised words, no digits/punctuation)
#     words = low.split()
#     if 2 <= len(words) <= 4 and all(w.isalpha() for w in words):
#         return True
#     return False


# def file_hash(content: bytes) -> str:
#     return hashlib.sha256(content).hexdigest()[:16]


# # ── Text extractors ───────────────────────────────────────────────────────────
# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     import io
#     text = ""
#     try:
#         with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text()
#                 if t:
#                     text += t + "\n"
#         if text.strip():
#             return text
#     except Exception:
#         pass
#     try:
#         doc = fitz.open(stream=file_bytes, filetype="pdf")
#         for page in doc:
#             text += page.get_text() + "\n"
#     except Exception:
#         pass
#     return text


# def extract_text_from_docx(file_bytes: bytes) -> str:
#     import io
#     doc = Document(io.BytesIO(file_bytes))
#     return "\n".join(p.text for p in doc.paragraphs)


# def extract_text_from_txt(file_bytes: bytes) -> str:
#     for enc in ("utf-8", "latin-1", "cp1252"):
#         try:
#             return file_bytes.decode(enc)
#         except Exception:
#             continue
#     return ""


# def extract_text(file_bytes: bytes, filename: str) -> str:
#     ext = Path(filename).suffix.lower()
#     if ext == ".pdf":
#         return extract_text_from_pdf(file_bytes)
#     elif ext in (".docx", ".doc"):
#         return extract_text_from_docx(file_bytes)
#     elif ext == ".txt":
#         return extract_text_from_txt(file_bytes)
#     return ""


# # ── Heuristic helpers ─────────────────────────────────────────────────────────
# def extract_email(text: str) -> str:
#     for token in text.split():
#         token = token.strip("(),;<>\"'\n\r")
#         if "@" in token and "." in token.split("@")[-1] and len(token) > 5:
#             return token
#     return ""


# def extract_name(text: str) -> str:
#     lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
#     for line in lines[:6]:
#         if "@" in line or len(line) > 50:
#             continue
#         words = line.split()
#         if 2 <= len(words) <= 5:
#             if all(ch.isalpha() or ch in " .-'" for ch in line):
#                 return line
#     return lines[0][:40] if lines else "Unknown"


# def extract_experience_years(text: str, sections: dict = None) -> float:
#     import datetime
#     current_year = datetime.datetime.now().year
#     ranges = []
#     if sections is None:
#         sections = split_sections(text)
#     edu_text  = sections.get("education", "")
#     work_text = text.replace(edu_text, " ") if edu_text.strip() else text

#     year_pattern = re.compile(
#         r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
#         r'((?:19|20)\d{2})'
#         r'\s*(?:-|–|—|to)\s*'
#         r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
#         r'((?:19|20)\d{2}|present|current|now|till date|date)',
#         re.IGNORECASE
#     )

#     for match in year_pattern.finditer(work_text):
#         s       = int(match.group(2))
#         end_raw = match.group(4).strip().lower()
#         e       = current_year if end_raw in ("present", "current", "now", "till date", "date") else int(end_raw)
#         if not (1980 <= s <= current_year and s <= e <= current_year + 1):
#             continue
#         ranges.append((s, e))

#     if ranges:
#         ranges.sort()
#         merged = [list(ranges[0])]
#         for start, end in ranges[1:]:
#             if start <= merged[-1][1]:
#                 merged[-1][1] = max(merged[-1][1], end)
#             else:
#                 merged.append([start, end])
#         total = sum(e - s for s, e in merged)
#         return round(min(total, 40.0), 1)

#     fallback_pattern = re.compile(
#         r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?(?:\s+experience)?',
#         re.IGNORECASE
#     )
#     matches = fallback_pattern.findall(work_text)
#     if matches:
#         return round(min(max(int(m) for m in matches), 40.0), 1)
#     return 0.0


# def extract_skills_from_text(text: str, skill_sections_text: str = "") -> list:
#     """
#     KEY FIX: Extract skills ONLY from skill-relevant sections.

#     Priority order:
#       1. If we have dedicated skill/experience/project section text → use ONLY that.
#       2. Otherwise fall back to full text BUT skip personal-info lines.

#     This prevents name tokens, address tokens, and email prefixes from
#     becoming "skills" that then get fuzzy-matched against JD terms.

#     No hardcoded taxonomy — LLM does semantic matching at query time.
#     """
#     skills = []
#     seen   = set()

#     # Use skill-section text if available, else fall back to full text
#     source_text = skill_sections_text.strip() if skill_sections_text.strip() else text

#     for line in source_text.split("\n"):
#         line = line.strip()
#         if len(line) < 3 or len(line) > 120:
#             continue

#         # ── NEW: skip personal info lines entirely ────────────────────────────
#         if _is_personal_info_line(line):
#             continue

#         cleaned = ""
#         for ch in line:
#             if ch.isalnum() or ch in " .+#/-_":
#                 cleaned += ch
#             else:
#                 cleaned += " "

#         raw_tokens = [t.strip(".-_/") for t in cleaned.split() if t.strip(".-_/")]

#         for i, tok in enumerate(raw_tokens):
#             tok_lower = tok.lower()
#             # ── NEW: skip very short tokens and pure-digit tokens ─────────────
#             if len(tok_lower) < 3 or tok_lower in STOP_WORDS or tok_lower.isdigit():
#                 continue
#             if tok_lower not in seen:
#                 seen.add(tok_lower)
#                 skills.append(tok_lower)

#             # 2-gram — only if first word is not a stop-word
#             if i + 1 < len(raw_tokens):
#                 next_tok = raw_tokens[i + 1].lower()
#                 if next_tok in STOP_WORDS or next_tok.isdigit():
#                     continue
#                 bigram = (tok_lower + " " + next_tok).strip()
#                 if bigram not in seen and len(bigram) <= 40:
#                     parts = bigram.split()
#                     if not any(p in STOP_WORDS for p in parts):
#                         seen.add(bigram)
#                         skills.append(bigram)

#     return skills[:200]


# # ── Section splitter ──────────────────────────────────────────────────────────
# def split_sections(text: str) -> dict:
#     lines    = text.split("\n")
#     sections = {k: [] for k in ["full", "skills", "experience", "projects",
#                                  "education", "certifications", "summary"]}
#     current  = "full"

#     for line in lines:
#         stripped = line.strip().lower().rstrip(":").strip()
#         matched  = None
#         for section, headers in SECTION_HEADERS.items():
#             if stripped in headers:
#                 matched = section
#                 break
#         if matched:
#             current = matched
#         else:
#             sections[current].append(line)

#     return {k: "\n".join(v).strip() for k, v in sections.items()}


# # ── Main pipeline (0 API calls) ───────────────────────────────────────────────
# def parse_resume(file_bytes: bytes, filename: str) -> dict:
#     raw_text = extract_text(file_bytes, filename)
#     if not raw_text.strip():
#         return {}

#     sections  = split_sections(raw_text)
#     exp_years = extract_experience_years(raw_text, sections)

#     # ── FIX: build skill-section source text (exclude header + contact area) ──
#     # Combine skill-relevant sections; fall back to full text only if empty
#     skill_source_parts = [
#         sections.get("skills", ""),
#         sections.get("experience", ""),
#         sections.get("projects", ""),
#         sections.get("summary", ""),
#         sections.get("certifications", ""),
#     ]
#     skill_source = "\n".join(p for p in skill_source_parts if p.strip())

#     skills = extract_skills_from_text(raw_text, skill_sections_text=skill_source)

#     chunk_labels = {
#         "skills":         sections.get("skills") or " ".join(skills[:60]),
#         "experience":     sections.get("experience") or "",
#         "projects":       sections.get("projects") or "",
#         "education":      sections.get("education") or "",
#         "certifications": sections.get("certifications") or "",
#         "summary":        sections.get("summary") or "",
#         "full":           raw_text[:3000],
#     }
#     chunks = {label: text.strip() for label, text in chunk_labels.items() if text.strip()}

#     return {
#         "filename":         filename,
#         "file_hash":        file_hash(file_bytes),
#         "name":             extract_name(raw_text),
#         "email":            extract_email(raw_text),
#         "skills":           skills,
#         "experience_years": exp_years,
#         "raw_text":         raw_text,
#         "chunks":           chunks,
#     }


# yeh final tha














# """
# Resume Parser — Zero API calls at upload time.
# ADDS: verified_skills (proven in experience/projects) vs claimed_skills (skills section only)
# """

# import hashlib
# import re
# import pdfplumber
# import fitz
# from pathlib import Path
# from docx import Document


# SECTION_HEADERS = {
#     "experience":     {
#         "work experience", "professional experience", "employment history",
#         "experience", "work history", "career history",
#     },
#     "education":      {
#         "education", "academic background", "qualifications",
#         "degrees", "academic qualifications",
#     },
#     "skills":         {
#         "skills", "technical skills", "core competencies", "technologies",
#         "tech stack", "expertise", "key skills", "competencies",
#         "tools & technologies", "tools and technologies",
#         "tools", "software", "platforms", "tech tools", "applications",
#     },
#     "projects":       {
#         "projects", "personal projects", "key projects",
#         "portfolio", "selected projects", "project experience",
#     },
#     "certifications": {
#         "certifications", "certificates", "awards",
#         "achievements", "licenses", "credentials",
#     },
#     "summary":        {
#         "summary", "objective", "profile", "about me",
#         "overview", "professional summary", "career objective",
#         "executive summary",
#     },
# }

# STOP_WORDS = {
#     "and", "or", "the", "a", "an", "in", "on", "at", "to", "for",
#     "of", "with", "is", "are", "was", "were", "be", "been", "have",
#     "has", "had", "do", "did", "will", "would", "could", "should",
#     "may", "might", "i", "me", "my", "we", "our", "you", "your",
#     "he", "she", "it", "they", "their", "this", "that", "these",
#     "those", "not", "no", "nor", "but", "so", "yet", "both", "either",
#     "as", "if", "then", "than", "by", "from", "up", "about", "into",
#     "through", "during", "including", "used", "using", "worked",
#     "work", "working", "experience", "strong", "good", "knowledge",
#     "understanding", "ability", "skills", "skill", "years", "year",
#     "month", "months", "present", "current", "responsible", "team",
#     "also", "well", "new", "use", "develop", "development",
#     "mr", "ms", "dr", "sir", "dear", "hello", "hi",
#     "street", "road", "lane", "avenue", "ave", "blvd", "st", "rd",
#     "city", "state", "country", "zip", "postcode", "uk", "us", "usa",
#     "phone", "mobile", "tel", "email", "linkedin", "github", "portfolio",
#     "page", "cv", "resume",
# }

# _PERSONAL_INFO_PREFIXES = (
#     "email", "phone", "mobile", "tel", "address", "linkedin",
#     "github", "twitter", "website", "portfolio", "location",
#     "nationality", "date of birth", "dob", "gender",
# )


# def _is_personal_info_line(line: str) -> bool:
#     low = line.lower().strip()
#     if any(low.startswith(p) for p in _PERSONAL_INFO_PREFIXES):
#         return True
#     if re.match(r"^\d+\s+\w", low):
#         return True
#     words = low.split()
#     if 2 <= len(words) <= 4 and all(w.isalpha() for w in words):
#         return True
#     return False


# def file_hash(content: bytes) -> str:
#     return hashlib.sha256(content).hexdigest()[:16]


# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     import io
#     text = ""
#     try:
#         with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text()
#                 if t:
#                     text += t + "\n"
#         if text.strip():
#             return text
#     except Exception:
#         pass
#     try:
#         doc = fitz.open(stream=file_bytes, filetype="pdf")
#         for page in doc:
#             text += page.get_text() + "\n"
#     except Exception:
#         pass
#     return text


# def extract_text_from_docx(file_bytes: bytes) -> str:
#     import io
#     doc = Document(io.BytesIO(file_bytes))
#     return "\n".join(p.text for p in doc.paragraphs)


# def extract_text_from_txt(file_bytes: bytes) -> str:
#     for enc in ("utf-8", "latin-1", "cp1252"):
#         try:
#             return file_bytes.decode(enc)
#         except Exception:
#             continue
#     return ""


# def extract_text(file_bytes: bytes, filename: str) -> str:
#     ext = Path(filename).suffix.lower()
#     if ext == ".pdf":
#         return extract_text_from_pdf(file_bytes)
#     elif ext in (".docx", ".doc"):
#         return extract_text_from_docx(file_bytes)
#     elif ext == ".txt":
#         return extract_text_from_txt(file_bytes)
#     return ""


# def extract_email(text: str) -> str:
#     for token in text.split():
#         token = token.strip("(),;<>\"'\n\r")
#         if "@" in token and "." in token.split("@")[-1] and len(token) > 5:
#             return token
#     return ""


# def extract_name(text: str) -> str:
#     lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
#     for line in lines[:6]:
#         if "@" in line or len(line) > 50:
#             continue
#         words = line.split()
#         if 2 <= len(words) <= 5:
#             if all(ch.isalpha() or ch in " .-'" for ch in line):
#                 return line
#     return lines[0][:40] if lines else "Unknown"


# def extract_experience_years(text: str, sections: dict = None) -> float:
#     import datetime
#     current_year = datetime.datetime.now().year
#     ranges = []
#     if sections is None:
#         sections = split_sections(text)
#     edu_text  = sections.get("education", "")
#     work_text = text.replace(edu_text, " ") if edu_text.strip() else text

#     year_pattern = re.compile(
#         r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
#         r'((?:19|20)\d{2})'
#         r'\s*(?:-|–|—|to)\s*'
#         r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
#         r'((?:19|20)\d{2}|present|current|now|till date|date)',
#         re.IGNORECASE
#     )
#     for match in year_pattern.finditer(work_text):
#         s       = int(match.group(2))
#         end_raw = match.group(4).strip().lower()
#         e       = current_year if end_raw in ("present", "current", "now", "till date", "date") else int(end_raw)
#         if not (1980 <= s <= current_year and s <= e <= current_year + 1):
#             continue
#         ranges.append((s, e))

#     if ranges:
#         ranges.sort()
#         merged = [list(ranges[0])]
#         for start, end in ranges[1:]:
#             if start <= merged[-1][1]:
#                 merged[-1][1] = max(merged[-1][1], end)
#             else:
#                 merged.append([start, end])
#         total = sum(e - s for s, e in merged)
#         return round(min(total, 40.0), 1)

#     fallback_pattern = re.compile(
#         r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?(?:\s+experience)?',
#         re.IGNORECASE
#     )
#     matches = fallback_pattern.findall(work_text)
#     if matches:
#         return round(min(max(int(m) for m in matches), 40.0), 1)
#     return 0.0


# def extract_skills_from_text(text: str, skill_sections_text: str = "") -> list:
#     skills = []
#     seen   = set()
#     source_text = skill_sections_text.strip() if skill_sections_text.strip() else text

#     for line in source_text.split("\n"):
#         line = line.strip()
#         if len(line) < 3 or len(line) > 120:
#             continue
#         if _is_personal_info_line(line):
#             continue

#         cleaned = ""
#         for ch in line:
#             if ch.isalnum() or ch in " .+#/-_":
#                 cleaned += ch
#             else:
#                 cleaned += " "

#         raw_tokens = [t.strip(".-_/") for t in cleaned.split() if t.strip(".-_/")]

#         for i, tok in enumerate(raw_tokens):
#             tok_lower = tok.lower()
#             if len(tok_lower) < 3 or tok_lower in STOP_WORDS or tok_lower.isdigit():
#                 continue
#             if tok_lower not in seen:
#                 seen.add(tok_lower)
#                 skills.append(tok_lower)

#             if i + 1 < len(raw_tokens):
#                 next_tok = raw_tokens[i + 1].lower()
#                 if next_tok in STOP_WORDS or next_tok.isdigit():
#                     continue
#                 bigram = (tok_lower + " " + next_tok).strip()
#                 if bigram not in seen and len(bigram) <= 40:
#                     parts = bigram.split()
#                     if not any(p in STOP_WORDS for p in parts):
#                         seen.add(bigram)
#                         skills.append(bigram)

#     return skills[:200]


# def split_sections(text: str) -> dict:
#     lines    = text.split("\n")
#     sections = {k: [] for k in ["full", "skills", "experience", "projects",
#                                  "education", "certifications", "summary"]}
#     current  = "full"

#     for line in lines:
#         stripped = line.strip().lower().rstrip(":").strip()
#         matched  = None
#         for section, headers in SECTION_HEADERS.items():
#             if stripped in headers:
#                 matched = section
#                 break
#         if matched:
#             current = matched
#         else:
#             sections[current].append(line)

#     return {k: "\n".join(v).strip() for k, v in sections.items()}


# # ── Main pipeline ─────────────────────────────────────────────────────────────
# def parse_resume(file_bytes: bytes, filename: str) -> dict:
#     raw_text = extract_text(file_bytes, filename)
#     if not raw_text.strip():
#         return {}

#     sections  = split_sections(raw_text)
#     exp_years = extract_experience_years(raw_text, sections)

#     skills_section_text = sections.get("skills", "")
#     proof_section_text  = "\n".join(filter(None, [
#         sections.get("experience", ""),
#         sections.get("projects", ""),
#         sections.get("summary", ""),
#     ]))
#     full_source = "\n".join(filter(None, [
#         skills_section_text,
#         proof_section_text,
#         sections.get("certifications", ""),
#     ]))

#     skills = extract_skills_from_text(raw_text, skill_sections_text=full_source)

#     proof_tokens       = set(extract_skills_from_text(raw_text, skill_sections_text=proof_section_text))
#     skills_only_tokens = set(extract_skills_from_text(raw_text, skill_sections_text=skills_section_text))

#     # Raw lowercase of experience+projects+summary for direct substring check.
#     # Catches skills that section-splitter put under "skills" header but are
#     # actually mentioned inline in experience/project bullet points too.
#     proof_raw_lower = proof_section_text.lower()

#     verified_skills = []
#     claimed_skills  = []
#     for sk in skills:
#         if sk in proof_tokens:
#             verified_skills.append(sk)
#         elif sk.lower() in proof_raw_lower:
#             # Skill literally appears in experience/projects text → verified
#             verified_skills.append(sk)
#         elif sk in skills_only_tokens:
#             claimed_skills.append(sk)
#         else:
#             verified_skills.append(sk)  # certifications etc → verified

#     chunk_labels = {
#         "skills":         skills_section_text or " ".join(skills[:60]),
#         "experience":     sections.get("experience") or "",
#         "projects":       sections.get("projects") or "",
#         "education":      sections.get("education") or "",
#         "certifications": sections.get("certifications") or "",
#         "summary":        sections.get("summary") or "",
#         "full":           raw_text[:3000],
#     }
#     chunks = {label: text.strip() for label, text in chunk_labels.items() if text.strip()}

#     return {
#         "filename":         filename,
#         "file_hash":        file_hash(file_bytes),
#         "name":             extract_name(raw_text),
#         "email":            extract_email(raw_text),
#         "skills":           skills,
#         "verified_skills":  verified_skills,
#         "claimed_skills":   claimed_skills,
#         "experience_years": exp_years,
#         "raw_text":         raw_text,
#         "chunks":           chunks,
#     }



"""
Resume Parser — Zero API calls at upload time.

Text extraction  : pdfplumber / PyMuPDF / python-docx
Skill extraction : plain text tokenization — no LLM, no regex taxonomy,
                   no hardcoded skill list. Raw tokens become the skills
                   bag-of-words; the LLM in retrieval_engine does the
                   semantic understanding at query time.
Experience years : simple digit scan (no regex — just token inspection)
"""

import hashlib
import pdfplumber
import fitz  # PyMuPDF
from pathlib import Path
from docx import Document


# ── Section headers (plain string match — no regex) ───────────────────────────
SECTION_HEADERS = {
    "experience":     {
        "work experience", "professional experience", "employment history",
        "experience", "work history", "career history",
    },
    "education":      {
        "education", "academic background", "qualifications",
        "degrees", "academic qualifications",
    },
    "skills":         {
        "skills", "technical skills", "core competencies", "technologies",
        "tech stack", "expertise", "key skills", "competencies",
        "tools & technologies", "tools and technologies",
        "tools", "software", "platforms", "tech tools", "applications",
    },
    "projects":       {
        "projects", "personal projects", "key projects",
        "portfolio", "selected projects", "project experience",
    },
    "certifications": {
        "certifications", "certificates", "awards",
        "achievements", "licenses", "credentials",
    },
    "summary":        {
        "summary", "objective", "profile", "about me",
        "overview", "professional summary", "career objective",
        "executive summary",
    },
}

# Stop-words to drop from skill tokens (common English words, not skills)
STOP_WORDS = {
    "and", "or", "the", "a", "an", "in", "on", "at", "to", "for",
    "of", "with", "is", "are", "was", "were", "be", "been", "have",
    "has", "had", "do", "did", "will", "would", "could", "should",
    "may", "might", "i", "me", "my", "we", "our", "you", "your",
    "he", "she", "it", "they", "their", "this", "that", "these",
    "those", "not", "no", "nor", "but", "so", "yet", "both", "either",
    "as", "if", "then", "than", "by", "from", "up", "about", "into",
    "through", "during", "including", "used", "using", "worked",
    "work", "working", "experience", "strong", "good", "knowledge",
    "understanding", "ability", "skills", "skill", "years", "year",
    "month", "months", "present", "current", "responsible", "team",
    "also", "well", "new", "use", "develop", "development",
}


def file_hash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()[:16]


# ── Text extractors ───────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    import io
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        if text.strip():
            return text
    except Exception:
        pass
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
    except Exception:
        pass
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    import io
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(enc)
        except Exception:
            continue
    return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        return extract_text_from_txt(file_bytes)
    return ""


# ── Heuristic helpers (no regex, no API) ──────────────────────────────────────

def extract_email(text: str) -> str:
    for token in text.split():
        token = token.strip("(),;<>\"'\n\r")
        if "@" in token and "." in token.split("@")[-1] and len(token) > 5:
            return token
    return ""


def extract_name(text: str) -> str:
    """First short line that looks like a person's name."""
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
    for line in lines[:6]:
        if "@" in line or len(line) > 50:
            continue
        words = line.split()
        if 2 <= len(words) <= 5:
            if all(ch.isalpha() or ch in " .-'" for ch in line):
                return line
    return lines[0][:40] if lines else "Unknown"


import re

# def extract_experience_years(text: str) -> float:
#     import datetime, re
#     current_year = datetime.datetime.now().year
#     ranges = []

#     year_pattern = re.compile(
#         r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
#         r'((?:19|20)\d{2})'
#         r'\s*(?:-|–|—|to)\s*'
#         r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
#         r'((?:19|20)\d{2}|present|current|now|till date|date)',
#         re.IGNORECASE
#     )

#     for match in year_pattern.finditer(text):
#         s = int(match.group(2))
#         end_raw = match.group(4).strip().lower()
#         e = current_year if end_raw in ("present", "current", "now", "till date", "date") else int(end_raw)

#         if not (1980 <= s <= current_year and s <= e <= current_year + 1):
#             continue
#         ranges.append((s, e))

#     if ranges:
#         # Merge overlapping ranges
#         ranges.sort()
#         merged = [list(ranges[0])]
#         for start, end in ranges[1:]:
#             if start <= merged[-1][1]:
#                 merged[-1][1] = max(merged[-1][1], end)
#             else:
#                 merged.append([start, end])
#         total = sum(e - s for s, e in merged)
#         return round(min(total, 40.0), 1)

#     # Fallback — "X years", "X+ years", "X yrs" pattern
#     fallback_pattern = re.compile(
#         r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?(?:\s+experience)?',
#         re.IGNORECASE
#     )
#     matches = fallback_pattern.findall(text)
#     if matches:
#         return round(min(max(int(m) for m in matches), 40.0), 1)

#     return 0.0


def extract_experience_years(text: str, sections: dict = None) -> float:
    import datetime, re
    current_year = datetime.datetime.now().year
    ranges = []
    if sections is None:
        sections = split_sections(text)
    edu_text = sections.get("education", "")
    work_text = text.replace(edu_text, " ") if edu_text.strip() else text

    year_pattern = re.compile(
        r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
        r'((?:19|20)\d{2})'
        r'\s*(?:-|–|—|to)\s*'
        r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*)?'
        r'((?:19|20)\d{2}|present|current|now|till date|date)',
        re.IGNORECASE
    )

    for match in year_pattern.finditer(work_text):
        s = int(match.group(2))
        end_raw = match.group(4).strip().lower()
        e = current_year if end_raw in ("present", "current", "now", "till date", "date") else int(end_raw)

        if not (1980 <= s <= current_year and s <= e <= current_year + 1):
            continue
        ranges.append((s, e))

    if ranges:
        ranges.sort()
        merged = [list(ranges[0])]
        for start, end in ranges[1:]:
            if start <= merged[-1][1]:
                merged[-1][1] = max(merged[-1][1], end)
            else:
                merged.append([start, end])
        total = sum(e - s for s, e in merged)
        return round(min(total, 40.0), 1)

    # Fallback
    fallback_pattern = re.compile(
        r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?(?:\s+experience)?',
        re.IGNORECASE
    )
    matches = fallback_pattern.findall(work_text)
    if matches:
        return round(min(max(int(m) for m in matches), 40.0), 1)

    return 0.0

def extract_skills_from_text(text: str) -> list:
    """
    Extract meaningful multi-word and single-word tokens from resume text.
    Strategy:
      1. Split into lines → extract noun phrases (consecutive capitalized or
         lowercase meaningful words) from bullet points / skill sections.
      2. Deduplicate, filter stop-words, keep tokens 2-40 chars.

    No hardcoded taxonomy — LLM does semantic matching at query time.
    This just gives us a fast keyword bag for skill-overlap scoring.
    """
    skills = []
    seen   = set()

    for line in text.split("\n"):
        line = line.strip()
        # Skip very short or very long lines (headers, paragraphs)
        if len(line) < 3 or len(line) > 120:
            continue

        # Clean punctuation but keep meaningful chars
        cleaned = ""
        for ch in line:
            if ch.isalnum() or ch in " .+#/-_":
                cleaned += ch
            else:
                cleaned += " "

        # Split into tokens
        raw_tokens = [t.strip(".-_/") for t in cleaned.split()]
        raw_tokens = [t for t in raw_tokens if t]

        # Sliding window: 1-gram and 2-gram skill phrases
        for i, tok in enumerate(raw_tokens):
            tok_lower = tok.lower()
            if len(tok_lower) < 2 or tok_lower in STOP_WORDS:
                continue
            if tok_lower not in seen:
                seen.add(tok_lower)
                skills.append(tok_lower)

            # 2-gram
            if i + 1 < len(raw_tokens):
                bigram = (tok_lower + " " + raw_tokens[i+1].lower()).strip()
                if bigram not in seen and len(bigram) <= 40:
                    # Only keep bigrams where neither part is a stop-word
                    parts = bigram.split()
                    if not any(p in STOP_WORDS for p in parts):
                        seen.add(bigram)
                        skills.append(bigram)

    return skills[:200]  # cap at 200 — more than enough


# ── Section splitter ──────────────────────────────────────────────────────────

def split_sections(text: str) -> dict:
    lines    = text.split("\n")
    sections = {k: [] for k in ["full", "skills", "experience", "projects",
                                 "education", "certifications", "summary"]}
    current  = "full"

    for line in lines:
        stripped = line.strip().lower().rstrip(":").strip()
        matched  = None
        for section, headers in SECTION_HEADERS.items():
            if stripped in headers:
                matched = section
                break
        if matched:
            current = matched
        else:
            sections[current].append(line)

    return {k: "\n".join(v).strip() for k, v in sections.items()}


# ── Main pipeline (0 API calls) ───────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Full parse pipeline — no API calls.
    Skills bag + section chunks are used by retrieval_engine for:
      - BM25 keyword matching
      - Skill-overlap scoring (against LLM-extracted JD skills)
      - Dense embedding (full text chunks)
    """
    raw_text = extract_text(file_bytes, filename)
    if not raw_text.strip():
        return {}

    sections  = split_sections(raw_text)
    skills    = extract_skills_from_text(raw_text)
    exp_years = extract_experience_years(raw_text, sections)

    chunk_labels = {
        "skills":         sections.get("skills") or " ".join(skills[:60]),
        "experience":     sections.get("experience") or "",
        "projects":       sections.get("projects") or "",
        "education":      sections.get("education") or "",
        "certifications": sections.get("certifications") or "",
        "summary":        sections.get("summary") or "",
        "full":           raw_text[:3000],
    }
    chunks = {label: text.strip() for label, text in chunk_labels.items() if text.strip()}

    return {
        "filename":         filename,
        "file_hash":        file_hash(file_bytes),
        "name":             extract_name(raw_text),
        "email":            extract_email(raw_text),
        "skills":           skills,
        "experience_years": exp_years,
        "raw_text":         raw_text,
        "chunks":           chunks,
    }